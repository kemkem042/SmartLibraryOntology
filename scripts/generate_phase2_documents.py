from __future__ import annotations

import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
REPORT_TEMPLATE = Path(
    "/Users/kemalmertceyhan/Downloads/KnowledgeEngineeringandOntologies_20252026_ProjectReport-1.docx"
)


def clear_document(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def add_paragraph_with_fallback(doc: Document, text: str, style: str | None = None) -> None:
    if style:
        try:
            doc.add_paragraph(text, style=style)
            return
        except KeyError:
            pass
    doc.add_paragraph(text)


def add_markdown_to_doc(doc: Document, markdown_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            continue

        if stripped.startswith("### "):
            add_paragraph_with_fallback(doc, stripped[4:], "Heading 3")
            continue

        if stripped.startswith("## "):
            add_paragraph_with_fallback(doc, stripped[3:], "Heading 2")
            continue

        if stripped.startswith("# "):
            add_paragraph_with_fallback(doc, stripped[2:], "Title")
            continue

        if re.match(r"^\d+\.\s", stripped):
            add_paragraph_with_fallback(doc, stripped)
            continue

        if stripped.startswith("- "):
            add_paragraph_with_fallback(doc, f"- {stripped[2:]}")
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            add_paragraph_with_fallback(doc, stripped)
            continue

        add_paragraph_with_fallback(doc, stripped)


def build_report() -> None:
    doc = Document(str(REPORT_TEMPLATE))
    clear_document(doc)
    add_markdown_to_doc(doc, ROOT / "docs/reports/project-report-v2.md")
    output = ROOT / "docs/reports/SmartLibraryOntology_Project_Report_v2.docx"
    doc.save(str(output))


def build_specification() -> None:
    doc = Document()
    add_markdown_to_doc(doc, ROOT / "docs/specification-v2.md")
    output = ROOT / "docs/Ontology_Requirements_Specification_v2.docx"
    doc.save(str(output))


if __name__ == "__main__":
    build_report()
    build_specification()
